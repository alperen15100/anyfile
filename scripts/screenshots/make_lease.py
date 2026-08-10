#!/usr/bin/env python3
"""
Builds the sample tenancy agreement as a .docx and converts it to PDF with
LibreOffice.

Two things this file works around, both found on Android System WebView 133
with the pdf.js Gander vendors:

1. A table near the top of a page blanks everything drawn above it on that
   page. The title, subtitle and running header all vanished while the table
   and the body below it rendered. Probe paragraphs inserted above the title
   vanished as well, and the blank band grew to match, so it is not a fixed
   clip. Ghostscript renders the same file correctly. The parties block is
   therefore laid out as labelled paragraphs rather than a table.
2. reportlab output hit the same thing and lost the page-1 footer as well, so
   the document is built as .docx and converted with LibreOffice.

Every name, address and figure is invented.
"""
import subprocess
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
OUT = HERE / "out"
OUT.mkdir(exist_ok=True)
SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

NAVY = "1F3A5F"
GREY = "6E6E6E"
SERIF = "Times New Roman"


def _field(paragraph, instr):
    """Insert a Word field (used for PAGE / NUMPAGES in the footer)."""
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run._r.addnext(fld)


def _fixed_grid(table, widths_cm):
    """Pin the column widths.

    Setting cell.width alone is not enough: Word and LibreOffice both fall back
    to the autofit algorithm unless the layout is declared fixed and the
    tblGrid carries the same widths.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    grid = tbl.find(qn("w:tblGrid"))
    for col in list(grid):
        grid.remove(col)
    for cm in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(cm * 567)))
        grid.append(gc)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def build_docx(path):
    d = docx.Document()

    sec = d.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    normal = d.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)

    # Running header, right aligned.
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("REF: AC-14/2026-08")
    hr.font.size = Pt(8)
    hr.font.name = SERIF
    hr.font.color.rgb = RGBColor.from_string("909090")

    # Footer: description on the left, page x of y on the right.
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run("Assured Shorthold Tenancy — 14 Alder Court, Kestrel Bay KB7 2QN\t\tPage ")
    fr.font.size = Pt(8)
    fr.font.name = SERIF
    fr.font.color.rgb = RGBColor.from_string("909090")
    _field(fp, "PAGE")
    tail = fp.add_run(" of ")
    tail.font.size = Pt(8)
    tail.font.name = SERIF
    tail.font.color.rgb = RGBColor.from_string("909090")
    _field(fp, "NUMPAGES")

    def para(text, size=10.5, bold=False, italic=False, color=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None, after=7, before=0):
        p = d.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_after = Pt(after)
        pf.space_before = Pt(before)
        if indent:
            pf.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.name = SERIF
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        return p

    para("RESIDENTIAL TENANCY AGREEMENT", size=20, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para("Assured Shorthold Tenancy under Part I, Chapter II of the Housing Act "
         "(fictional jurisdiction of Kestrel Bay)", size=9.5, italic=True,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    parties = [
        # No registration number. An invented 8-digit one follows the real Companies
        # House format closely enough that it will sometimes name an actual company,
        # and the first attempt did. See the README before inventing an identifier.
        ("Landlord",
         "Marisol Okonkwo-Reyes, of 3 Thornhill Row, Kestrel Bay KB4 8LT, acting "
         "through her managing agent, Aldergate Property Care."),
        ("Tenants",
         "Dev Raghunathan and Priya Raghunathan, jointly and severally liable for the "
         "whole of the rent and for every obligation in this agreement."),
        ("Property",
         "The whole of the two-bedroom first-floor flat known as 14 Alder Court, Marrow "
         "Lane, Kestrel Bay KB7 2QN, together with the fixtures listed in the Inventory "
         "at Schedule 2 and the allocated parking bay numbered 14."),
        ("Term",
         "Twelve (12) months certain, beginning on 1 September 2026 and ending on "
         "31 August 2027, after which the tenancy continues as a statutory periodic "
         "tenancy from month to month until ended under clause 11."),
        ("Rent",
         "£1,340 per calendar month, payable in advance on the 1st day of each month by "
         "standing order to the client account named in Schedule 1. The first payment "
         "falls due on 1 September 2026."),
        ("Deposit",
         "£1,546, being five weeks' rent, protected within 30 days of receipt in an "
         "approved custodial scheme. The scheme reference is issued with the prescribed "
         "information at Schedule 3."),
    ]
    for label, body in parties:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(5)
        pf.left_indent = Cm(3.1)
        pf.first_line_indent = Cm(-3.1)
        lab = p.add_run(label + "\t")
        lab.bold = True
        lab.font.size = Pt(10)
        lab.font.name = SERIF
        txt = p.add_run(body)
        txt.font.size = Pt(10)
        txt.font.name = SERIF

    para("", after=4)
    para("It is agreed that the Landlord lets and the Tenants take the Property for the "
         "Term at the Rent, on the terms set out below. Where the Tenants are more than "
         "one person, every obligation binds each of them individually and all of them "
         "together. Words in the singular include the plural, and a reference to a clause "
         "is a reference to a clause of this agreement.")

    clauses = [
        ("Rent and outgoings", [
            "The Tenants shall pay the Rent without deduction or set-off. If any part of "
            "the Rent remains unpaid for fourteen days after it became due, interest "
            "accrues on the unpaid amount at three per cent above the Bank of Kestrel "
            "base rate, calculated daily from the date the payment fell due until it is "
            "paid in full, whether or not the Landlord has formally demanded it.",
            "The Tenants shall pay all charges for electricity, gas, water, wastewater, "
            "broadband and council tax relating to the Property during the Term, "
            "including any standing charges and any reconnection fee arising from a "
            "supply cut off because of non-payment, and shall provide meter readings on "
            "the first and last day of the Term.",
        ], []),
        ("Condition and repair", [
            "The Tenants shall keep the interior of the Property in the same clean and "
            "tenantable condition in which they received it, fair wear and tear excepted, "
            "and shall make good any damage caused by them, by anyone living with them or "
            "by any visitor. The Landlord remains responsible for the structure and "
            "exterior, for the installations for the supply of water, gas, electricity and "
            "sanitation, and for the space and water heating.",
        ], [
            "The Tenants shall report any disrepair, water ingress, or failure of a heating "
            "or sanitary installation to the managing agent within five working days of "
            "becoming aware of it. A defect that worsens because it was not reported may be "
            "charged to the Tenants to the extent of the additional cost.",
            "The Tenants shall test the smoke alarms monthly and shall not remove a battery "
            "from any alarm except to replace it immediately.",
            "The Tenants shall keep the Property adequately ventilated and heated so as to "
            "avoid condensation, and shall not dry laundry indoors without opening a window "
            "or running the extractor fan.",
            "The Tenants shall not redecorate, drill, or fix anything to a wall without the "
            "Landlord's written consent, which will not be unreasonably withheld for "
            "ordinary picture hooks.",
        ]),
        ("Use of the Property", [
            "The Tenants shall use the Property as a single private dwelling for themselves "
            "and for no more than two other permitted occupiers named in Schedule 1. The "
            "Tenants shall not carry on any trade or business at the Property, other than "
            "quiet office work that generates no visitors, no signage and no deliveries "
            "beyond ordinary domestic volume.",
            "The Tenants shall not do anything at the Property that is a nuisance or "
            "annoyance to neighbours, that invalidates the buildings insurance, or that is "
            "illegal or immoral. Noise audible outside the Property between 11pm and 7am is "
            "treated as a breach of this clause.",
        ], []),
        ("Access by the Landlord", [
            "The Landlord and anyone authorised by the Landlord may enter the Property at "
            "reasonable hours on at least twenty-four hours' written notice to inspect its "
            "condition, to carry out repairs, or to show it to a prospective tenant or "
            "purchaser during the final two months of the Term. In a genuine emergency, "
            "including fire, flood, or a suspected gas leak, entry may be made without "
            "notice and the Landlord shall tell the Tenants what happened as soon as "
            "practicable afterwards.",
        ], []),
        ("Assignment and subletting", [
            "The Tenants shall not assign, underlet, charge, or part with or share "
            "possession of the whole or any part of the Property. Taking in a lodger, "
            "listing the Property on a short-let platform, or allowing a person not named "
            "in Schedule 1 to reside at the Property for more than fourteen consecutive "
            "nights each amount to a breach of this clause.",
        ], []),
        ("Pets", [
            "The Tenants shall not keep any animal, bird, reptile or insect at the Property "
            "without the Landlord's prior written consent. Consent given for one animal does "
            "not extend to another. Where consent is given, the Tenants shall have the "
            "carpets and soft furnishings professionally cleaned at the end of the Term and "
            "shall produce the receipt.",
        ], []),
        ("Insurance", [
            "The Landlord shall insure the building and the Landlord's fixtures and "
            "fittings. That policy does not cover the Tenants' own possessions, and the "
            "Tenants are strongly advised to take out their own contents insurance. The "
            "Tenants shall not keep anything at the Property that would increase the premium "
            "or make the policy void, and shall pay any increase in premium caused by "
            "something they do.",
        ], []),
        ("Alterations and installations", [
            "The Tenants shall not alter the Property, install a satellite dish or an "
            "external aerial, change any lock, or add a smart doorbell or camera that "
            "records a shared hallway, without the Landlord's written consent. Where consent "
            "is given, the Tenants shall reinstate the Property at the end of the Term if "
            "the Landlord asks them to.",
        ], []),
        ("Garden and communal areas", [
            "The Tenants shall keep the allocated parking bay clear of anything other than a "
            "roadworthy and taxed vehicle, and shall not store a bicycle, pushchair or any "
            "other item in the communal stairwell, which is a protected fire escape route. "
            "Items left in the stairwell may be removed without notice and the cost of "
            "removal charged to the Tenants.",
        ], []),
        ("Deposit and deductions", [
            "At the end of the Term the Landlord may propose deductions from the Deposit for "
            "unpaid rent, for damage beyond fair wear and tear, for cleaning to bring the "
            "Property back to the standard recorded in the Inventory, and for the reasonable "
            "cost of removing anything the Tenants leave behind. The Landlord shall put any "
            "proposed deduction to the Tenants in writing with supporting evidence within "
            "ten working days of the end of the Term, and any part of the Deposit not in "
            "dispute shall be returned within the same period.",
        ], [
            "A deduction may not be made for an item that the Inventory records as already "
            "damaged or missing at the start of the Term.",
            "Where the Tenants dispute a proposed deduction, either party may refer the "
            "dispute to the free adjudication service operated by the deposit scheme, and "
            "both parties agree to be bound by its decision.",
        ]),
        ("Ending the tenancy", [
            "Either party may end the periodic tenancy that follows the Term by giving the "
            "other written notice expiring on the last day of a rental period, being at "
            "least one month's notice from the Tenants and at least two months' notice from "
            "the Landlord. Notice given by the Tenants must be given by all of them.",
            "On the last day of the Term or of any periodic tenancy the Tenants shall give "
            "up the Property with vacant possession, return every key, fob and remote "
            "control issued to them, and provide a forwarding address.",
        ], []),
        ("Notices", [
            "A notice under this agreement is validly given if it is delivered by hand, sent "
            "by first-class post, or sent to the email address each party has given in "
            "Schedule 1. A posted notice is treated as received on the second working day "
            "after posting, and an emailed notice on the next working day after it is sent, "
            "unless the sender receives a delivery failure message.",
        ], []),
    ]

    for i, (head, bodies, subs) in enumerate(clauses, start=1):
        para(f"{i}. {head}", size=11.5, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=9, after=4)
        for b in bodies:
            para(b)
        for j, s in enumerate(subs, start=1):
            para(f"{i}.{j}   {s}", indent=0.7, after=5)

    para("", after=10)
    para("Signed by the Landlord's agent", size=10, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=2, before=6)
    para("_______________________________     Aldergate Property Care"
         "     Date: ______________", size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=10)
    para("Signed by the Tenants", size=10, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    para("_______________________________     Dev Raghunathan"
         "     Date: ______________", size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=6)
    para("_______________________________     Priya Raghunathan"
         "     Date: ______________", size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=6)

    d.save(path)
    return path


def to_pdf(docx_path):
    subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf",
                    "--outdir", str(OUT), str(docx_path)],
                   check=True, capture_output=True)
    return OUT / (Path(docx_path).stem + ".pdf")


if __name__ == "__main__":
    tmp = HERE / "lease-src.docx"
    build_docx(tmp)
    pdf = to_pdf(tmp)
    final = OUT / "Tenancy Agreement - 14 Alder Court.pdf"
    if pdf != final:
        pdf.replace(final)
    print("wrote", final, final.stat().st_size // 1024, "KB")
